import subprocess
import sys
from pathlib import Path
import pytest

SCRIPT = Path(__file__).parents[1] / "skills/aem/edge-delivery-services/aem-doc-converter/scripts/md-to-docx.py"


def run_script(input_path, output_path):
    return subprocess.run(
        [sys.executable, str(SCRIPT), str(input_path), str(output_path)],
        capture_output=True, text=True
    )


def test_produces_docx_file(simple_md_path, tmp_path):
    out = tmp_path / "out.docx"
    result = run_script(simple_md_path, out)
    assert result.returncode == 0
    assert out.exists()
    assert out.stat().st_size > 0


def test_docx_contains_headings(simple_md_path, tmp_path):
    from docx import Document
    out = tmp_path / "out.docx"
    run_script(simple_md_path, out)
    doc = Document(str(out))
    styles = [p.style.name for p in doc.paragraphs]
    assert any("Heading" in s for s in styles)


def test_docx_contains_text(simple_md_path, tmp_path):
    from docx import Document
    out = tmp_path / "out.docx"
    run_script(simple_md_path, out)
    doc = Document(str(out))
    full_text = " ".join(p.text for p in doc.paragraphs)
    assert "Section One" in full_text
    assert "Section Two" in full_text


def test_eds_blog_tables_preserved(eds_blog_md_path, tmp_path):
    from docx import Document
    out = tmp_path / "out.docx"
    run_script(eds_blog_md_path, out)
    doc = Document(str(out))
    assert len(doc.tables) > 0


def test_missing_input_exits_nonzero(tmp_path):
    out = tmp_path / "out.docx"
    result = run_script(tmp_path / "nonexistent.md", out)
    assert result.returncode != 0
