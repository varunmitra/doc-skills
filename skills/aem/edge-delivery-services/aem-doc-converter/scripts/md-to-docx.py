#!/usr/bin/env python3
"""
md-to-docx.py — Convert Markdown to Word document (.docx).

Preserves: headings (H1-H4), paragraphs, bold, italic, tables,
fenced code blocks, bullet/numbered lists.
Note: uses python-docx locally — zero LLM tokens.

Usage: python md-to-docx.py <input.md> <output.docx>
"""
import sys
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    print("Error: Missing dependency. Run: pip install python-docx", file=sys.stderr)
    sys.exit(1)


def add_table_from_eds(doc, table_lines):
    """Parse EDS +-delimited table text and add as docx table."""
    rows = []
    current_row = []
    for line in table_lines:
        stripped = line.strip()
        if stripped.startswith("+") and set(stripped) <= set("+-=|"):
            if current_row:
                rows.append([c.strip() for c in current_row])
                current_row = []
        elif stripped.startswith("|"):
            cells = [c.strip() for c in stripped.split("|")[1:-1]]
            if current_row:
                for i, c in enumerate(cells):
                    if i < len(current_row):
                        if c:
                            current_row[i] = (current_row[i] + " " + c).strip()
                    else:
                        current_row.append(c)
            else:
                current_row = cells
    if current_row:
        rows.append([c.strip() for c in current_row])
    if not rows:
        return

    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            if c_idx < col_count:
                clean = re.sub(r'\*\*(.+?)\*\*', r'\1', cell_text)
                clean = re.sub(r'\*(.+?)\*', r'\1', clean)
                clean = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', clean)
                cell = table.rows[r_idx].cells[c_idx]
                cell.text = clean
                if r_idx == 0:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True


def add_inline_runs(para, text):
    """Add text with bold/italic inline markdown to a paragraph."""
    remaining = text
    while remaining:
        bold_m = re.match(r'\*\*(.+?)\*\*', remaining, re.DOTALL)
        italic_m = re.match(r'\*(.+?)\*', remaining, re.DOTALL)
        code_m = re.match(r'`(.+?)`', remaining)

        candidates = [(m, kind) for m, kind in [
            (bold_m, 'bold'), (italic_m, 'italic'), (code_m, 'code')
        ] if m]

        if not candidates:
            para.add_run(remaining)
            break

        first = min(candidates, key=lambda x: x[0].start())
        m, kind = first

        if m.start() > 0:
            para.add_run(remaining[:m.start()])

        run = para.add_run(m.group(1))
        if kind == 'bold':
            run.bold = True
        elif kind == 'italic':
            run.italic = True
        elif kind == 'code':
            run.font.name = "Courier New"
            run.font.size = Pt(9)

        remaining = remaining[m.end():]


def convert(input_path, output_path):
    source = Path(input_path).read_text(encoding="utf-8")
    doc = Document()

    lines = source.splitlines()
    i = 0
    in_table = False
    table_lines_buf = []
    in_code = False
    code_lines = []

    def flush_table():
        nonlocal in_table, table_lines_buf
        if table_lines_buf:
            add_table_from_eds(doc, table_lines_buf)
        in_table = False
        table_lines_buf = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # fenced code block
        if stripped.startswith("```"):
            if in_table:
                flush_table()
            if not in_code:
                in_code = True
                code_lines = []
            else:
                in_code = False
                para = doc.add_paragraph()
                run = para.add_run("\n".join(code_lines))
                run.font.name = "Courier New"
                run.font.size = Pt(9)
                code_lines = []
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # EDS table lines
        if stripped.startswith("+") and re.match(r'^\+[-=+|]+', stripped):
            in_table = True
            table_lines_buf.append(line)
            i += 1
            continue
        if in_table and (stripped.startswith("|") or stripped.startswith("+")):
            table_lines_buf.append(line)
            i += 1
            continue
        if in_table:
            flush_table()

        # headings
        h_match = re.match(r'^(#{1,4})\s+(.*)', stripped)
        if h_match:
            level = len(h_match.group(1))
            doc.add_heading(h_match.group(2), level=level)
            i += 1
            continue

        # horizontal rule
        if re.match(r'^---+$', stripped):
            doc.add_paragraph("─" * 40)
            i += 1
            continue

        # bullet list
        bullet_m = re.match(r'^[-*]\s+(.*)', stripped)
        if bullet_m:
            para = doc.add_paragraph(style="List Bullet")
            add_inline_runs(para, bullet_m.group(1))
            i += 1
            continue

        # numbered list
        num_m = re.match(r'^\d+\.\s+(.*)', stripped)
        if num_m:
            para = doc.add_paragraph(style="List Number")
            add_inline_runs(para, num_m.group(1))
            i += 1
            continue

        # blank line
        if not stripped:
            i += 1
            continue

        # regular paragraph
        para = doc.add_paragraph()
        add_inline_runs(para, stripped)
        i += 1

    if in_table:
        flush_table()

    doc.save(str(output_path))


def main():
    if len(sys.argv) != 3:
        print("Usage: md-to-docx.py <input.md> <output.docx>", file=sys.stderr)
        sys.exit(1)
    input_path = Path(sys.argv[1])
    if not input_path.exists():
        print(f"Error: Input file not found: {input_path}", file=sys.stderr)
        sys.exit(1)
    convert(input_path, sys.argv[2])
    print(f"✓ Converted to {sys.argv[2]}")


if __name__ == "__main__":
    main()
